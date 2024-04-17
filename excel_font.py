import pandas as pd
from docx import Document
from docx.shared import Pt  # For font size

excel_file = r'C:\Users\njlao\Desktop\Requirement.xlsx'
sheet_name = 'Sheet1'  # Or your specific sheet name
document = Document()  # Create a new Document

# Reading the Excel file
df = pd.read_excel(excel_file, sheet_name=sheet_name)

def add_requirement_to_document(document, row):
    Number = row['number'] if not pd.isna(row['name']) else ''
    Name = row['name'] if not pd.isna(row['name']) else ''
    ShallShould = row['shall/should'] if not pd.isna(row['shall/should']) else ''
    Description = row['description'] if not pd.isna(row['description']) else ''
    Comment = row['comment'] if not pd.isna(row['comment']) else ''
    
    # Add Name in bold and Arial font
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{Number} {Name} {ShallShould} {Description}\n\n")
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)  # Example font size
    
    # Additional details with formatting applied as needed
    paragraph.add_run("Requirement Type: ").bold = True
    paragraph.add_run("Functional\n\n").font.name = 'Arial'
    
    paragraph.add_run("ID: ").bold = True
    paragraph.add_run(f"{Name}\n\n").font.name = 'Arial'
    
    paragraph.add_run("Description\n").bold = True
    if Comment:
        # Split the comments by semicolon and add each as a new line
        for comment in Comment.split(';'):
            paragraph.add_run(f"{comment.strip()}\n").font.name = 'Arial'
    else:
        paragraph.add_run("\n").font.name = 'Arial'
    
    # Add other static information similarly...
    paragraph.add_run("Change Information: ").bold = True
    paragraph.add_run("No change issue detected.\n\n").font.name = 'Arial'
    
    paragraph.add_run("Implementation Status\n ").bold = True
    paragraph.add_run("Total: 1, Implemented: 0, Justified: 0, None: 1\n\n").font.name = 'Arial'

    paragraph.add_run("Verification Status\n").bold = True
    paragraph.add_run("Total: 1, Passed: 0, Justified: 0, Failed: 0, Unexecuted: 0, None: 1\n\n").font.name = 'Arial'
    # Continue with the rest of the static info
    # ...

# Iterate through each row in the DataFrame and format the requirement
for index, row in df.iterrows():
    add_requirement_to_document(document, row)

# Save the document
document.save(r'C:\Users\njlao\Desktop\requirements.docx')